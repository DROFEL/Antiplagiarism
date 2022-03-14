from unittest import case
from soupsieve import match
from docx import Document
import sys

#droppedFile = sys.argv[1]
#print(droppedFile)

document = Document("/home/drofel/Desktop/Projects/Antiplagiarism/doc.docx")

for count, paragraph in enumerate(document.paragraphs):

    #strings are immutable so to modify it I converted it to a list
    text = list(paragraph.text)
    
    #loop for detecting and raplacing characters
    for i, chr in enumerate(text):
        if(chr == "a"):
            text[i] = "а"
        elif(chr == "p"):
            text[i] = "р"
        elif(chr == "c"):
            text[i] = "с"
        elif(chr == "x"):
            text[i] = "х"
        elif(chr == "y"):
            text[i] = "у"
        elif(chr == "e"):
            text[i] = "е"

    #convert list back to string
    paragraph.text = "".join(text)
    document.paragraphs[count] = paragraph
#save In new document    
document.save("/home/drofel/Desktop/Projects/Antiplagiarism/docEdit.docx")

#*for i, chr in enumerate(text):
#        if(chr == "a"):
#            text[i] = "а"
#        elif(chr == "p"):
#            text[i] = "р"
#        elif(chr == "c"):
#            text[i] = "с"
#        elif(chr == "x"):
#            text[i] = "х"
#        elif(chr == "y"):
#            text[i] = "у"
#        elif(chr == "e"):
#            text[i] = "е"